#!/usr/bin/env python3
"""
CASP Agent Scaffold

Interactive command-line agent that collects CASP data, saves a JSON session,
and renders a Word (.docx) document using a Word template for text placement
and python-docx for table insertion.
"""

from __future__ import annotations

import copy
import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent
SCHEMA_PATH = BASE_DIR / "casp_schema.json"
EXAMPLE_PATH = BASE_DIR / "sample_casp_data.json"
DEFAULT_OUTPUT = BASE_DIR / "generated_casp.docx"
TEMPLATE_PATH = BASE_DIR / "CASPTemplate.docx"

DEFAULT_AIMS = [
    "Detail the duties and responsibilities of the staff",
    "Detail the implementation of control measures from the SST and Risk Assessment",
    "Detail the administrative and support requirements",
]
DEFAULT_DOCUMENTS = [
    "Risk Assessment",
    "Training Programme",
    "Stores List",
    "Exercise and Action Safety Plan (EASP)",
    "MOD Form 1930 Safe Activity Assurance Form (SAAF)",
    "Joining Instructions"
]


def load_schema() -> Dict[str, Any]:
    return json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))


def save_json(path: Path, data: Dict[str, Any]) -> None:
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def input_text(prompt: str, default: str = "") -> str:
    if default:
        keep = input(f"{prompt}: current value is '{default}'. Keep it? [Y/n]: ").strip().lower()
        if keep in ("", "y", "yes"):
            return default
        return input(f"Enter new value for {prompt}: ").strip()
    return input(f"{prompt}: ").strip()


def input_list(prompt: str, default: Optional[List[str]] = None) -> List[str]:
    default = default or []

    if default:
        print(f"{prompt}: current values are:")
        for item in default:
            print(f"  - {item}")
        keep = input("Keep these values? [Y/n]: ").strip().lower()
        if keep in ("", "y", "yes"):
            return default

    print(f"Enter new values for {prompt} (blank line to finish)")
    items: List[str] = []
    while True:
        raw = input("  - ").strip()
        if not raw:
            break
        items.append(raw)

    return items


def input_yes_no(prompt: str, default: bool = False) -> bool:
    default_text = "Y/n" if default else "y/N"
    raw = input(f"{prompt} [{default_text}]: ").strip().lower()
    if not raw:
        return default
    return raw in {"y", "yes", "true", "1"}


def input_person(role: str, default: Optional[Dict[str, str]] = None) -> Dict[str, str]:
    default = default or {}
    print(f"\nEnter details for {role}")
    return {
        "Role": role,
        "Rank": input_text("  Rank", default.get("Rank", "")),
        "Name": input_text("  Name", default.get("Name", "")),
        "Qualification": input_text("  Qualification", default.get("Qualification", "")),
        "Remarks": input_text("  Remarks", default.get("Remarks", "")),
    }


def collect_repeating_people(role: str, current: Optional[List[Dict[str, str]]] = None) -> List[Dict[str, str]]:
    collected: List[Dict[str, str]] = []
    current = current or []

    if current and input_yes_no(f"Reuse existing {role} entries from loaded data", True):
        collected.extend(current)

    while input_yes_no(f"Add another {role}"):
        collected.append(input_person(role))

    return collected


def collect_basic_details(data: Dict[str, Any]) -> None:
    print("\n=== Unit details ===")
    unit = data["Unit"]
    unit["Appointment"] = input_text("Unit appointment", unit["Appointment"])
    unit["Name"] = input_text("Unit title", unit["Name"])
    unit["Address"]["Line1"] = input_text("Address line 1", unit["Address"]["Line1"])
    unit["Address"]["Line2"] = input_text("Address line 2", unit["Address"]["Line2"])
    unit["Address"]["County"] = input_text("County", unit["Address"]["County"])
    unit["Address"]["Postcode"] = input_text("Post code", unit["Address"]["Postcode"])
    unit["Phone"] = input_text("Phone", unit["Phone"])
    unit["Email"] = input_text("Email", unit["Email"])

    print("\n=== Activity details ===")
    activity = data["Activity"]
    activity["Code"] = input_text("Activity code", activity["Code"])
    activity["Name"] = input_text("Activity name", activity["Name"])
    activity["Type"] = input_text("Activity type", activity["Type"])
    activity["SubUnit"] = input_text("Sub-unit", activity["SubUnit"])
    activity["PrimaryLocation"] = input_text("Primary location", activity["PrimaryLocation"])
    activity["StartDate"] = input_text("Start date/time", activity["StartDate"])
    activity["EndDate"] = input_text("End date/time", activity["EndDate"])
    activity["DocumentDate"] = input_text("Document date", activity["DocumentDate"])
    activity["Purpose"] = input_text("Purpose / reasons for activity", activity["Purpose"])
    activity["CFAVCount"] = input_text("Approx CFAV count", str(activity["CFAVCount"]))
    activity["CadetCount"] = input_text("Approx cadet count", str(activity["CadetCount"]))
    activity["ConsultedWith"] = input_text("Consulted with", activity["ConsultedWith"])

    print("\n=== Risk assessments ===")
    ra = data["RiskAssessments"]
    ra["CoreTrainingAtRFCAorDIO"]["Date"] = input_text(
        "Core training at RFCA/DIO locations RA date",
        ra["CoreTrainingAtRFCAorDIO"]["Date"],
    )
    ra["AF5010C"]["Reference"] = input_text("AF-5010C reference", ra["AF5010C"]["Reference"])
    ra["AF5010C"]["Date"] = input_text("AF-5010C date", ra["AF5010C"]["Date"])
    data["Authority"]["AtCApprover"] = input_text("AtC approver", data["Authority"]["AtCApprover"])

    print("\n=== Supporting documents ===")
    data["Documents"] = input_list("List additional documents", DEFAULT_DOCUMENTS)

    print("\n=== Aims of this document ===")
    data["Aims"] = input_list("Enter aim statements", DEFAULT_AIMS)


def collect_staff(data: Dict[str, Any]) -> None:
    print("\n=== Core staff roles ===")
    existing = {entry["Role"]: entry for entry in data.get("Staff", [])}
    roles = [
        "Senior Activity Owner",
        "Activity Owner",
        "Activity Planner",
        "Personnel Owner",
    ]
    data["Staff"] = [input_person(role, existing.get(role, {})) for role in roles]

    print("\n=== Sub-Activity Owners ===")
    existing_sub = data.get("SubActivityOwners", [])
    data["SubActivityOwners"] = []

    sub_index = 0
    while input_yes_no("Do you want to add a Sub-Activity Owner?", False):
        default_sub = existing_sub[sub_index] if sub_index < len(existing_sub) else {}
        data["SubActivityOwners"].append(input_person("Sub-Activity Owner", default_sub))
        sub_index += 1

    print("\n=== First Aiders ===")
    existing_first = data.get("FirstAiders", [])
    data["FirstAiders"] = []

    default_first = existing_first[0] if existing_first else {}
    data["FirstAiders"].append(input_person("First Aider", default_first))

    first_index = 1
    while input_yes_no("Do you want to add another First Aider?", False):
        default_first = existing_first[first_index] if first_index < len(existing_first) else {}
        data["FirstAiders"].append(input_person("First Aider", default_first))
        first_index += 1

    print("\n=== Drivers ===")
    existing_drivers = data.get("Drivers", [])
    data["Drivers"] = []

    default_driver = existing_drivers[0] if existing_drivers else {}
    data["Drivers"].append(input_person("Driver", default_driver))

    driver_index = 1
    while input_yes_no("Do you want to add another Driver?", False):
        default_driver = existing_drivers[driver_index] if driver_index < len(existing_drivers) else {}
        data["Drivers"].append(input_person("Driver", default_driver))
        driver_index += 1



def collect_locations(data: Dict[str, Any]) -> None:
    print("\n=== Activity locations ===")
    locations: List[Dict[str, str]] = []
    existing = data.get("Locations", [])
    reuse = existing if existing and input_yes_no("Reuse existing locations from loaded data", True) else []
    locations.extend(reuse)

    if not reuse and input_yes_no("Add first activity location", True):
        locations.append({
            "Name": input_text("  Location name"),
            "RiskControls": input_text("  Risk controls"),
        })

    while input_yes_no("Add another activity location", False):
        locations.append({
            "Name": input_text("  Location name"),
            "RiskControls": input_text("  Risk controls"),
        })

    data["Locations"] = locations

    print("\n=== Medical references per location ===")
    medical_locations: List[Dict[str, Any]] = []
    for loc in data["Locations"]:
        print(f"\nMedical details for {loc['Name']}")
        medical_locations.append({
            "Name": loc["Name"],
            "What3Words": input_text("  What3Words"),
            "SafetyVehicle": {
                "Location": input_text("  Safety vehicle location"),
                "Remarks": input_text("  Safety vehicle remarks"),
            },
            "UrgentCare": {
                "Location": input_text("  Urgent care / walk-in location"),
                "Remarks": input_text("  Urgent care remarks"),
            },
            "AandE": {
                "Location": input_text("  A&E location"),
                "Remarks": input_text("  A&E remarks"),
            },
            "AmbulanceRV": {
                "Location": input_text("  Ambulance RV location"),
                "Remarks": input_text("  Ambulance RV remarks"),
            },
        })

    data["MedicalLocations"] = medical_locations


def collect_distribution(data: Dict[str, Any]) -> None:
    print("\n=== Distribution ===")
    distribution = data["Distribution"]
    distribution["Distribution"] = input_list(
        "Distribution list",
        distribution.get("Distribution", []),
    )
    distribution["Info"] = input_list(
        "Info list",
        distribution.get("Info", []),
    )

    print("\n=== Signatory ===")
    signatory = data["Signatory"]
    signatory["Name"] = input_text("Signatory name", signatory["Name"])
    signatory["Rank"] = input_text("Signatory rank", signatory["Rank"])
    signatory["Appointment"] = input_text("Signatory appointment", signatory["Appointment"])



def remove_numbering(paragraph):
    pPr = paragraph._p.get_or_add_pPr()

    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)

def build_staff_table(doc, staff_list):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.style = "Table Grid"

    # column widths
    table.columns[0].width = Inches(0.5)  # Ser
    equal_width = Inches(1.50)
    table.columns[1].width = equal_width
    table.columns[2].width = equal_width
    table.columns[3].width = equal_width
    table.columns[4].width = equal_width
    # FORCE column 0 width on all rows (Word requires cell-level control)
    for row in table.rows:
        row.cells[0].width = Inches(0.5)

    # header row
    headers = ["Ser", "Role", "Name", "Qualification", "Remarks"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # data rows
    for idx, person in enumerate(staff_list, start=1):
        row_cells = table.add_row().cells
        row_cells[0].width = Inches(0.4)

        values = [
            str(idx),
            person.get("Role", ""),
            f"{person.get('Rank', '')} {person.get('Name', '')}".strip(),
            person.get("Qualification", ""),
            person.get("Remarks", ""),
        ]

        for i, val in enumerate(values):
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]

            

            p.style = doc.styles["Normal"]
            
            run = p.add_run(val)
            run.font.size = Pt(10)


    return table
def build_locations_table(doc, locations_list):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.style = "Table Grid"

    # column widths
    table.columns[0].width = Inches(0.04)   # Ser - slightly wider than the word "Ser"
    equal_width = Inches(3.0)              # Location and Risk Controls equal width
    table.columns[1].width = equal_width
    table.columns[2].width = equal_width
    # FORCE column 0 width on all rows (Word requires cell-level control)
    for row in table.rows:
        row.cells[0].width = Inches(0.4)


    # headers
    headers = ["Ser", "Location", "Risk Controls"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = doc.styles["Normal"]
        run = p.add_run(h)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        run.font.size = Pt(10)

    # data rows
    for idx, loc in enumerate(locations_list, start=1):
        row_cells = table.add_row().cells
        row_cells[0].width = Inches(0.4)

        values = [
            str(idx),
            loc.get("Name", ""),
            loc.get("RiskControls", ""),
        ]

        for i, val in enumerate(values):
            cell = row_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]

            remove_numbering(p)

            p.style = doc.styles["Normal"]
            run = p.add_run(val)
            run.font.size = Pt(10)

    return table

def insert_table_after(paragraph, table):
    paragraph._p.addnext(table._tbl)
 
    return table

def build_medical_locations_table(doc, medical_locations):
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.style = "Table Grid"

    # column widths
    table.columns[0].width = Inches(0.6)   # Ser
    equal_width = Inches(1.6)
    table.columns[1].width = equal_width
    table.columns[2].width = equal_width
    table.columns[3].width = equal_width
    table.columns[4].width = equal_width

    # force column 0 width
    for row in table.rows:
        row.cells[0].width = Inches(0.6)

    # headers
    headers = ["Ser", "Location", "Type", "Facility", "Remarks"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""

        p = cell.paragraphs[0]
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)

    # data
    ser = 1

    for loc in medical_locations:
        name = loc.get("Name", "")

        entries = [
            ("Safety Vehicle", loc.get("SafetyVehicle", {})),
            ("Urgent Care", loc.get("UrgentCare", {})),
            ("A&E", loc.get("AandE", {})),
            ("Ambulance RV", loc.get("AmbulanceRV", {})),
        ]

        for entry_type, entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].width = Inches(0.6)

            values = [
                str(ser),
                name,
                entry_type,
                entry.get("Location", ""),
                entry.get("Remarks", ""),
            ]

            for i, val in enumerate(values):
                cell = row_cells[i]
                cell.text = ""

                p = cell.paragraphs[0]
                remove_numbering(p)
                p.style = doc.styles["Normal"]

                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = p.add_run(val)
                run.font.size = Pt(10)

            ser += 1

    return table

def replace_paragraph_with_locations_table(doc, placeholder_text, locations_list):
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            table = build_locations_table(doc, locations_list)

            paragraph._p.addnext(table._tbl)
            paragraph._element.getparent().remove(paragraph._element)
            return

    print(f"Placeholder '{placeholder_text}' not found in document.")

def replace_paragraph_with_staff_table(doc, placeholder_text, staff_list):
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            table = build_staff_table(doc, staff_list)

            insert_table_after(paragraph, table)

            paragraph._element.getparent().remove(paragraph._element)
            return

    print(f"Placeholder '{placeholder_text}' not found in document.")

def replace_paragraph_with_medical_locations_table(doc, placeholder_text, medical_locations):
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:

            table = build_medical_locations_table(doc, medical_locations)

            paragraph._p.addnext(table._tbl)
            paragraph._element.getparent().remove(paragraph._element)
            return

    print(f"Placeholder '{placeholder_text}' not found in document.")

def render_docx(data: Dict[str, Any], output_path: Path = DEFAULT_OUTPUT) -> Path:
    # Pass 1: render normal template content with docxtpl
    tpl = DocxTemplate(str(TEMPLATE_PATH))

    context = {
        "Activity": data.get("Activity", {}),
        "Unit": data.get("Unit", {}),
        "RiskAssessments": data.get("RiskAssessments", {}),
        "Authority": data.get("Authority", {}),
        "Documents": data.get("Documents", []),
        "Aims": data.get("Aims", []),
        "Distribution": data.get("Distribution", {}),
        "Signatory": data.get("Signatory", {}),
        "Annexes": data.get("Annexes", []),
    }

    tpl.render(context)
    tpl.save(str(output_path))

    # Pass 2: reopen with python-docx and insert the real staff table
    doc = Document(str(output_path))
    replace_paragraph_with_locations_table(doc, "__LOCATIONS_TABLE__", data.get("Locations", []))

    all_staff = (
        data.get("Staff", []) +
        data.get("SubActivityOwners", []) +
        data.get("FirstAiders", []) +
        data.get("Drivers", [])
    )

    replace_paragraph_with_staff_table(doc, "__STAFF_TABLE__", all_staff)

   
# --- Medical Locations table ---
    replace_paragraph_with_medical_locations_table(
        doc,
        "__MEDICAL_LOCATIONS_TABLE__",
        data.get("MedicalLocations", [])
    )



    doc.save(str(output_path))

    return output_path


def choose_starting_data():
    base = load_schema()

    if input_yes_no("Load the included sample data", False):
        data = json.loads(EXAMPLE_PATH.read_text(encoding="utf-8"))
        data["Aims"] = data.get("Aims", DEFAULT_AIMS)
        data["Documents"] = data.get("Documents", DEFAULT_DOCUMENTS)
        data["Staff"] = data.get("Staff", [])

        if input_yes_no("Generate document now?", True):
            return data, True

        return data, False

    custom_path_raw = input_text("Path to existing JSON to load (leave blank for empty schema)")
    if custom_path_raw:
        custom_path = Path(custom_path_raw).expanduser().resolve()
        data = json.loads(custom_path.read_text(encoding="utf-8"))
        return data, False

    return copy.deepcopy(base), False


def main() -> None:
    print("CASP Agent Scaffold")
    print("===================")

    data, generate_now = choose_starting_data()

    if not generate_now:
        collect_basic_details(data)
        collect_staff(data)
        collect_locations(data)
        collect_distribution(data)

        data["Aims"] = input_list("Enter aim statements", DEFAULT_AIMS)
        data["Documents"] = input_list("List additional documents", DEFAULT_DOCUMENTS)

        json_output = Path(
            input_text("JSON output path", str(BASE_DIR / "casp_session.json"))
        ).expanduser().resolve()
        save_json(json_output, data)
        print(f"Saved session JSON to: {json_output}")

        docx_output = Path(
            input_text("DOCX output path", str(DEFAULT_OUTPUT))
        ).expanduser().resolve()
    else:
        docx_output = Path(DEFAULT_OUTPUT)

    render_docx(data, docx_output)
    print(f"Generated CASP DOCX: {docx_output}")


if __name__ == "__main__":
    main()